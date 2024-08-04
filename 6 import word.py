# pip install python-docx

import pandas as pd
from docx import Document

filename = "w_sample.docx"

document = Document(filename)

tables = document.tables

my_document = '''
              <!DOCTYPE html>
              <html>
              <head>
                  <title>Average salsry since 1991y Russia</title>
              </head>
              <body>   
              '''

for idx, table in enumerate(tables):

    my_table = "<table id='table_" + str(idx) + "'>"

    for row_idx, row in enumerate(table.rows):

        my_row = "<tr>"

        for cell in row.cells:
            my_cell = ""
            for paragraph in cell.paragraphs:
                my_cell = my_cell + paragraph.text

            if row_idx == 0:
                my_row = my_row + "<th>" + my_cell + "</th>"
            else:
                my_row = my_row + "<td>" + my_cell + "</td>"
        my_row = my_row + "</tr>"

        my_table = my_table + my_row
    my_table = my_table + "</table>\n"

    my_document = my_document + my_table

my_document = my_document + '</table></body></html>'

with open('word_tables.html', 'w') as file:
    file.write(my_document)

table_html = pd.read_html(my_document,
                          header=1)
data = table_html[0]

data = data.drop(['I', 'II', 'III', 'IV'], 1)